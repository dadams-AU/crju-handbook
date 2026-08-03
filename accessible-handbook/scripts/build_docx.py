#!/usr/bin/env python3
"""Build and post-process the faculty-editable handbook DOCX."""

from __future__ import annotations

import argparse
import copy
import re
import subprocess
import tempfile
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CSUF_BLUE = "00244E"
HEADER_FILL = "EAF2F8"
PORTRAIT_WIDTH = Inches(8.5)
PORTRAIT_HEIGHT = Inches(11)
LANDSCAPE_WIDTH = Inches(11)
LANDSCAPE_HEIGHT = Inches(8.5)


def set_style_font(style, name: str, size: float, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    style.element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def configure_reference(path: Path) -> None:
    doc = Document(path)
    section = doc.sections[0]
    section.page_width = PORTRAIT_WIDTH
    section.page_height = PORTRAIT_HEIGHT
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = doc.styles
    for style_name in ("Normal", "Body Text", "Compact", "First Paragraph"):
        if style_name in styles:
            set_style_font(styles[style_name], "Arial", 11)
    if "Title" in styles:
        set_style_font(styles["Title"], "Arial", 24, True)
    if "Subtitle" in styles:
        set_style_font(styles["Subtitle"], "Arial", 14)
    if "Author" in styles:
        set_style_font(styles["Author"], "Arial", 11)
    if "Date" in styles:
        set_style_font(styles["Date"], "Arial", 11)
    if "Heading 1" in styles:
        set_style_font(styles["Heading 1"], "Arial", 16, True)
        styles["Heading 1"].font.color.rgb = RGBColor(0x00, 0x24, 0x4E)
    if "Heading 2" in styles:
        set_style_font(styles["Heading 2"], "Arial", 13, True)
        styles["Heading 2"].font.color.rgb = RGBColor(0x00, 0x24, 0x4E)
    if "Heading 3" in styles:
        set_style_font(styles["Heading 3"], "Arial", 11, True)
    if "Hyperlink" in styles:
        styles["Hyperlink"].font.color.rgb = RGBColor(0x00, 0x5E, 0xA8)
        styles["Hyperlink"].font.color.theme_color = None
        styles["Hyperlink"].font.underline = True

    doc.save(path)


def section_properties(landscape: bool) -> OxmlElement:
    sect = OxmlElement("w:sectPr")
    sect_type = OxmlElement("w:type")
    sect_type.set(qn("w:val"), "nextPage")
    sect.append(sect_type)

    page_size = OxmlElement("w:pgSz")
    if landscape:
        page_size.set(qn("w:w"), str(LANDSCAPE_WIDTH.twips))
        page_size.set(qn("w:h"), str(LANDSCAPE_HEIGHT.twips))
        page_size.set(qn("w:orient"), "landscape")
    else:
        page_size.set(qn("w:w"), str(PORTRAIT_WIDTH.twips))
        page_size.set(qn("w:h"), str(PORTRAIT_HEIGHT.twips))
    sect.append(page_size)

    margins = OxmlElement("w:pgMar")
    margin = Inches(0.75 if landscape else 1).twips
    for side in ("top", "right", "bottom", "left"):
        margins.set(qn(f"w:{side}"), str(margin))
    margins.set(qn("w:header"), str(Inches(0.35).twips))
    margins.set(qn("w:footer"), str(Inches(0.35).twips))
    margins.set(qn("w:gutter"), "0")
    sect.append(margins)
    return sect


def insert_section_break_before(paragraph, landscape_for_previous: bool) -> None:
    break_paragraph = OxmlElement("w:p")
    properties = OxmlElement("w:pPr")
    properties.append(section_properties(landscape_for_previous))
    break_paragraph.append(properties)
    paragraph._p.addprevious(break_paragraph)


def set_final_portrait_section(doc: Document) -> None:
    body_sect = doc.element.body.sectPr
    replacement = section_properties(False)
    replacement.remove(replacement.find(qn("w:type")))
    body_sect.getparent().replace(body_sect, replacement)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "AAB7C4")


def scale_table_grid(table, total_width_twips: int) -> None:
    grid_columns = table._tbl.tblGrid.gridCol_lst
    widths = [int(col.get(qn("w:w")) or 1) for col in grid_columns]
    current_total = sum(widths) or len(widths)
    scaled = [max(1, round(width * total_width_twips / current_total)) for width in widths]
    for col, width in zip(grid_columns, scaled):
        col.set(qn("w:w"), str(width))

    factor = total_width_twips / current_total
    for row in table.rows:
        for cell in row.cells:
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is not None and tc_w.get(qn("w:type"), "dxa") == "dxa":
                old = int(tc_w.get(qn("w:w")) or 0)
                if old:
                    tc_w.set(qn("w:w"), str(max(1, round(old * factor))))


def copy_cell_content(source_cell, target_cell) -> None:
    """Copy rich cell content without carrying merge/split cell properties."""
    for child in list(target_cell._tc):
        if child.tag != qn("w:tcPr"):
            target_cell._tc.remove(child)
    for child in source_cell._tc:
        if child.tag != qn("w:tcPr"):
            target_cell._tc.append(copy.deepcopy(child))


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    for child in list(cell._tc):
        if child.tag != qn("w:tcPr"):
            cell._tc.remove(child)
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold


def replace_table(doc: Document, table_index: int, headers: list[str] | None = None) -> None:
    """Rebuild a table as a simple rectangular grid with no merged cells."""
    old_table = doc.tables[table_index]
    column_count = len(old_table.columns)
    new_table = doc.add_table(rows=0, cols=column_count)

    if headers:
        header_cells = new_table.add_row().cells
        for index, value in enumerate(headers):
            set_cell_text(header_cells[index], value, bold=True)

    for old_row in old_table.rows:
        new_cells = new_table.add_row().cells
        for index, new_cell in enumerate(new_cells):
            copy_cell_content(old_row.cells[index], new_cell)

    old_table._tbl.addprevious(new_table._tbl)
    old_table._tbl.getparent().remove(old_table._tbl)


def normalize_course_table(table) -> None:
    section_names = (
        "Core Curriculum",
        "Breadth Requirement",
        "Electives",
        "Related Fields Courses",
        "Upper Division Writing Requirement",
    )
    for row in table.rows[1:]:
        first_text = row.cells[0].text.strip()
        if any(first_text.startswith(name) for name in section_names):
            for cell in row.cells[1:]:
                set_cell_text(cell, "")


def remove_simulated_input_lines(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if re.fullmatch(r"[_\s]+", paragraph.text or "") and "_" in paragraph.text:
                        paragraph.clear()


def replace_paragraph_text(paragraph, text: str, bold_prefix: str | None = None) -> None:
    paragraph.clear()
    if bold_prefix and text.startswith(bold_prefix):
        paragraph.add_run(bold_prefix).bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def normalize_simulated_footnotes(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "****":
            paragraph._p.getparent().remove(paragraph._p)
            continue
        if not text.startswith("*"):
            continue

        if text.startswith("**Current"):
            clean = "Policy note: " + text.lstrip("* ")
            replace_paragraph_text(paragraph, clean, "Policy note: ")
        elif text.startswith("** Take") or text.startswith("* Take"):
            clean = "Transfer note: " + text.lstrip("* ")
            replace_paragraph_text(paragraph, clean, "Transfer note: ")
        elif text.startswith("* Most transfer"):
            clean = text.lstrip("* ").replace(" **Current", " Policy note: Current")
            clean = "Transfer note: " + clean
            replace_paragraph_text(paragraph, clean, "Transfer note: ")
        else:
            clean = "Note: " + text.lstrip("* ")
            replace_paragraph_text(paragraph, clean, "Note: ")


def descriptive_url_label(url: str) -> str:
    parsed = urlparse(url)
    host = parsed.netloc.lower().removeprefix("www.") or "website"
    segments = [unquote(part) for part in parsed.path.split("/") if part]
    generic = {"index", "default", "content", "page"}
    detail = ""
    for segment in reversed(segments):
        stem = re.sub(r"\.(aspx?|php|html?)$", "", segment, flags=re.I)
        stem = re.sub(r"[-_]+", " ", stem).strip()
        if stem and stem.lower() not in generic:
            detail = stem
            break
    if parsed.fragment:
        fragment = re.sub(r"[-_]+", " ", unquote(parsed.fragment)).strip()
        if fragment and fragment.lower() not in {"top", "main"}:
            detail = fragment
    return f"{host} — {detail}" if detail else f"{host} website"


def hyperlink_target(doc: Document, hyperlink) -> str:
    relationship_id = hyperlink.get(qn("r:id"))
    if relationship_id and relationship_id in doc.part.rels:
        return doc.part.rels[relationship_id].target_ref
    return ""


def normalize_hyperlinks(doc: Document) -> None:
    for hyperlink in doc.element.body.iter(qn("w:hyperlink")):
        target = hyperlink_target(doc, hyperlink)
        visible = "".join(node.text or "" for node in hyperlink.iter(qn("w:t"))).strip()
        if not target:
            continue

        comparison = visible.strip("()<> ").rstrip("/")
        target_comparison = target.strip("()<> ").rstrip("/")
        raw_url_text = comparison.startswith(("http://", "https://")) or comparison == target_comparison
        if raw_url_text and target.startswith(("http://", "https://")):
            visible = descriptive_url_label(target)
            text_nodes = list(hyperlink.iter(qn("w:t")))
            if text_nodes:
                text_nodes[0].text = visible
                for node in text_nodes[1:]:
                    node.text = ""

        if target.startswith("mailto:"):
            tooltip = f"Email {visible or target.removeprefix('mailto:')}"
        else:
            tooltip = f"Open {visible or descriptive_url_label(target)}"
        hyperlink.set(qn("w:tooltip"), tooltip[:250])

        for run in hyperlink.iter(qn("w:r")):
            run_properties = run.find(qn("w:rPr"))
            if run_properties is None:
                run_properties = OxmlElement("w:rPr")
                run.insert(0, run_properties)
            color = run_properties.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                run_properties.append(color)
            color.set(qn("w:val"), "005EA8")
            color.attrib.pop(qn("w:themeColor"), None)
            underline = run_properties.find(qn("w:u"))
            if underline is None:
                underline = OxmlElement("w:u")
                run_properties.append(underline)
            underline.set(qn("w:val"), "single")


def set_column_widths(table, widths_inches: list[float]) -> None:
    widths = [Inches(value).twips for value in widths_inches]
    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def format_table(table, index: int) -> None:
    table.style = "Table"
    set_table_borders(table)
    table.autofit = False
    matrix = 4 <= index <= 11
    font_size = 10 if matrix else 10.5

    if index == 13:
        set_column_widths(table, [5.0, 0.75, 0.75])
    elif matrix:
        scale_table_grid(table, Inches(9.5).twips)
    elif len(table.columns) >= 4:
        scale_table_grid(table, Inches(6.5).twips)

    if table.rows:
        header_row_count = 2 if matrix else 1
        for header_row in table.rows[:header_row_count]:
            set_repeat_table_header(header_row)
            for cell in header_row.cells:
                shade_cell(cell, HEADER_FILL)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0x00, 0x24, 0x4E)
                        run.font.bold = True

    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Arial")

    if index in (5, 7, 9, 11):
        # Give students enough room to hand-write plans on printed blank
        # worksheets without inflating their header and totals rows.
        for row in table.rows[2:-2]:
            row.height = Inches(0.55)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def add_page_numbers(doc: Document) -> None:
    for index, section in enumerate(doc.sections):
        section.footer.is_linked_to_previous = index != 0

    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def validate_accessibility_structure(doc: Document) -> None:
    merge_markers = sum(
        1
        for table in doc.tables
        for element in table._tbl.iter()
        if element.tag in (qn("w:gridSpan"), qn("w:vMerge"))
    )
    simulated_fields = sum(
        1
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if re.fullmatch(r"[_\s]+", paragraph.text or "") and "_" in paragraph.text
    )
    simulated_footnotes = [
        paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip().startswith("*")
    ]
    external_links = [
        hyperlink
        for hyperlink in doc.element.body.iter(qn("w:hyperlink"))
        if hyperlink_target(doc, hyperlink)
    ]
    missing_names = [link for link in external_links if not link.get(qn("w:tooltip"))]
    raw_url_labels = []
    low_contrast_links = []
    for hyperlink in external_links:
        visible = "".join(node.text or "" for node in hyperlink.iter(qn("w:t"))).strip()
        if visible.strip("()<> ").startswith(("http://", "https://")):
            raw_url_labels.append(visible)
        for run in hyperlink.iter(qn("w:r")):
            properties = run.find(qn("w:rPr"))
            color = properties.find(qn("w:color")) if properties is not None else None
            if color is None or color.get(qn("w:val")) != "005EA8":
                low_contrast_links.append(visible)

    failures = {
        "merged/split table cells": merge_markers,
        "simulated input lines": simulated_fields,
        "simulated footnotes": len(simulated_footnotes),
        "hyperlinks without names": len(missing_names),
        "URL-only hyperlink labels": len(raw_url_labels),
        "hyperlinks without accessible color": len(low_contrast_links),
    }
    active_failures = {name: count for name, count in failures.items() if count}
    if active_failures:
        raise RuntimeError(f"DOCX accessibility structure checks failed: {active_failures}")
    print("DOCX accessibility structure checks passed.")


def postprocess(path: Path) -> None:
    doc = Document(path)
    first_roadmap = None
    minor_heading = None

    in_roadmaps = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("SAMPLE COURSE MATRIX FOR STUDENTS") or text.startswith("SAMPLE MATRIX FOR TRANSFER STUDENTS"):
            in_roadmaps = True
            paragraph.paragraph_format.page_break_before = True
            if first_roadmap is None:
                first_roadmap = paragraph
        if text == "Blank Matrix for Planning Purposes":
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.keep_with_next = True
        if text == "CRIMINAL JUSTICE MINOR":
            minor_heading = paragraph
            in_roadmaps = False
        if in_roadmaps:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1

    if first_roadmap is None or minor_heading is None:
        raise RuntimeError("Could not locate road-map section boundaries in generated DOCX")

    insert_section_break_before(first_roadmap, landscape_for_previous=False)
    first_roadmap.paragraph_format.page_break_before = False
    insert_section_break_before(minor_heading, landscape_for_previous=True)
    set_final_portrait_section(doc)

    # LibreOffice flags every split/merged cell. Rebuild only the affected
    # tables as simple grids. Duplicate year labels are intentional: each
    # semester column then has a complete independent header.
    for table_index in (2, 4, 5, 6, 7, 8, 9, 10, 11):
        headers = None
        if table_index == 2:
            headers = ["Course code", "Course title", "Planned term", "Completed term", "Grade"]
        replace_table(doc, table_index, headers)
    normalize_course_table(doc.tables[2])
    replace_table(doc, 13)

    for index, table in enumerate(doc.tables):
        format_table(table, index)

    remove_simulated_input_lines(doc)
    normalize_simulated_footnotes(doc)
    normalize_hyperlinks(doc)

    # Refresh the section collection after inserting the section breaks, then
    # create one footer field inherited by the later sections.
    add_page_numbers(doc)

    validate_accessibility_structure(doc)
    doc.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--pandoc", default="pandoc")
    args = parser.parse_args()

    args.output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="crju-docx-") as temp_dir:
        temp = Path(temp_dir)
        reference = temp / "reference.docx"
        raw_output = temp / "handbook.docx"
        subprocess.run([args.pandoc, "/dev/null", "--from=markdown", "-o", str(reference)], check=True)
        configure_reference(reference)
        subprocess.run(
            [
                args.pandoc,
                str(args.source),
                "--from=markdown",
                f"--reference-doc={reference}",
                "--toc",
                "--toc-depth=2",
                "--resource-path=.:..",
                "-o",
                str(raw_output),
            ],
            check=True,
        )
        postprocess(raw_output)
        args.output.write_bytes(raw_output.read_bytes())


if __name__ == "__main__":
    main()
